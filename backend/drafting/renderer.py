"""조례 → DOCX/PDF 렌더러 (law-ebansimsa pipeline.export.renderer 포팅).

출력물은 drafting_sections에서 파생되며 DB에 저장하지 않는다 — 온디맨드 재생성.
"""
from __future__ import annotations

import io
from dataclasses import dataclass


@dataclass
class ExportSection:
    """출력용 조 단위 항목"""
    article_label: str  # "제1조(목적)"
    title: str          # "목적"
    body: str           # 조문 본문
    order: int


def _sorted(sections: list[ExportSection]) -> list[ExportSection]:
    return sorted(sections, key=lambda s: s.order)


def render_docx(project_title: str, municipality: str, sections: list[ExportSection]) -> bytes:
    """조례를 DOCX 바이트로 렌더한다."""
    from docx import Document

    doc = Document()
    doc.add_heading(project_title, level=0)
    if municipality:
        doc.add_paragraph(municipality)
    for sec in _sorted(sections):
        doc.add_heading(sec.article_label, level=1)
        doc.add_paragraph(sec.body)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_pdf(project_title: str, municipality: str, sections: list[ExportSection]) -> bytes:
    """조례를 PDF 바이트로 렌더한다."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    flow = [Paragraph(project_title, styles["Title"])]
    if municipality:
        flow.append(Paragraph(municipality, styles["Normal"]))
    flow.append(Spacer(1, 12))
    for sec in _sorted(sections):
        flow.append(Paragraph(sec.article_label, styles["Heading2"]))
        flow.append(Paragraph(sec.body, styles["Normal"]))
        flow.append(Spacer(1, 6))
    doc.build(flow)
    return buf.getvalue()
