"""
일괄 개정검토 처리 서비스
5단계 파이프라인: 대상선정 → 개정판별 → 제개정이유수집 → AI분석 → 보고서
"""
import logging
from datetime import datetime
from typing import Optional, AsyncGenerator

from sqlalchemy import select, func, and_, update
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload, joinedload

from backend.core.config import settings
from backend.models.batch_job import BatchJob, BatchJobItem
from backend.models.ordinance import Ordinance
from backend.models.law import Law
from backend.models.ordinance_law_mapping import OrdinanceLawMapping
from backend.models.law_revision_reason import LawRevisionReason
from backend.models.llm_analysis_result import LlmAnalysisResult
from backend.models.revision_detection_result import RevisionDetectionResult
from backend.services.revision_detection_service import RevisionDetectionService
from backend.services.llm_analysis_service import LlmAnalysisService, ConflictError

logger = logging.getLogger(__name__)


class BatchProcessingService:
    """일괄 개정검토 처리"""

    def __init__(self, db: AsyncSession):
        self.db = db

    # ===== Batch Job CRUD =====

    async def create_batch_job(
        self, name: str, filter_params: dict, user_id: Optional[int] = None
    ) -> BatchJob:
        """Step 1: 배치 작업 생성 + 필터 조건으로 대상 조례 선정"""
        job = BatchJob(
            name=name,
            created_by_id=user_id,
            current_step="step1_select",
            step1_status="completed",
            filter_params=filter_params,
        )
        self.db.add(job)
        await self.db.flush()

        # 필터 조건으로 조례 목록 조회
        query = select(Ordinance).where(Ordinance.status == "ACTIVE")

        if filter_params.get("category"):
            query = query.where(Ordinance.category == filter_params["category"])
        if filter_params.get("department"):
            dept = filter_params["department"]
            query = query.where(
                (Ordinance.department == dept) | (Ordinance.department.like(f"% {dept}"))
            )
        if filter_params.get("search"):
            query = query.where(Ordinance.name.ilike(f"%{filter_params['search']}%"))

        result = await self.db.execute(query.order_by(Ordinance.name))
        ordinances = list(result.scalars().all())

        # 각 조례에 대해 BatchJobItem 생성
        for ord in ordinances:
            # 자동 제외 로직
            step1_result = "included"
            step1_reason = None

            if ord.status == "ABOLISHED":
                step1_result = "excluded"
                step1_reason = "폐지"
            elif ord.no_parent_law:
                step1_result = "excluded"
                step1_reason = "상위법령없음 확인"

            # 상위법령 매핑 확인
            if step1_result == "included":
                mapping_check = await self.db.execute(
                    select(func.count()).where(
                        OrdinanceLawMapping.ordinance_id == ord.id
                    )
                )
                if mapping_check.scalar() == 0:
                    step1_result = "excluded"
                    step1_reason = "상위법령 미연결"

            item = BatchJobItem(
                batch_job_id=job.id,
                ordinance_id=ord.id,
                ordinance_name=ord.name,
                ordinance_department=ord.department,
                ordinance_category=ord.category,
                step1_result=step1_result,
                step1_reason=step1_reason,
                final_result="제외(" + step1_reason + ")" if step1_result == "excluded" else None,
            )
            self.db.add(item)

        await self.db.flush()
        await self.db.commit()
        return job

    async def get_batch_job(self, job_id: int) -> Optional[BatchJob]:
        result = await self.db.execute(
            select(BatchJob).where(BatchJob.id == job_id)
        )
        return result.scalar_one_or_none()

    async def get_batch_jobs(self) -> list[BatchJob]:
        result = await self.db.execute(
            select(BatchJob).order_by(BatchJob.created_at.desc())
        )
        return list(result.scalars().all())

    async def delete_batch_job(self, job_id: int) -> bool:
        job = await self.get_batch_job(job_id)
        if not job:
            return False
        await self.db.delete(job)
        await self.db.commit()
        return True

    async def get_batch_items(
        self,
        job_id: int,
        step_filter: Optional[str] = None,
        result_filter: Optional[str] = None,
        page: int = 1,
        size: int = 50,
    ) -> dict:
        """배치 항목 조회 (단계별 필터링, 페이지네이션)"""
        query = select(BatchJobItem).where(BatchJobItem.batch_job_id == job_id)

        if step_filter and result_filter:
            col = getattr(BatchJobItem, f"{step_filter}_result", None)
            if col is not None:
                query = query.where(col == result_filter)

        # 수동 제외 항목 필터
        count_query = select(func.count()).select_from(query.subquery())
        total = (await self.db.execute(count_query)).scalar() or 0

        query = query.order_by(BatchJobItem.ordinance_name)
        query = query.offset((page - 1) * size).limit(size)

        result = await self.db.execute(query)
        items = list(result.scalars().all())

        return {"items": items, "total": total, "page": page, "size": size}

    async def get_step_counts(self, job_id: int) -> dict:
        """각 단계별 결과 집계"""
        items_result = await self.db.execute(
            select(BatchJobItem).where(BatchJobItem.batch_job_id == job_id)
        )
        items = list(items_result.scalars().all())

        counts = {
            "total": len(items),
            "step1": {"included": 0, "excluded": 0},
            "step2": {"pending": 0, "needs_revision": 0, "no_revision": 0, "error": 0},
            "step3": {"pending": 0, "collected": 0, "no_data": 0, "error": 0},
            "step4": {"pending": 0, "needs_revision": 0, "no_revision": 0, "error": 0},
            "final": {},
        }

        for item in items:
            # Step 1
            counts["step1"][item.step1_result] = counts["step1"].get(item.step1_result, 0) + 1

            # Step 2
            if item.step2_result:
                counts["step2"][item.step2_result] = counts["step2"].get(item.step2_result, 0) + 1
            elif item.step1_result == "included":
                counts["step2"]["pending"] += 1

            # Step 3
            if item.step3_result:
                counts["step3"][item.step3_result] = counts["step3"].get(item.step3_result, 0) + 1
            elif item.step2_result == "needs_revision":
                counts["step3"]["pending"] += 1

            # Step 4
            if item.step4_result:
                counts["step4"][item.step4_result] = counts["step4"].get(item.step4_result, 0) + 1
            elif item.step3_result == "collected":
                counts["step4"]["pending"] += 1

            # Final
            if item.final_result:
                counts["final"][item.final_result] = counts["final"].get(item.final_result, 0) + 1

        return counts

    async def toggle_item_exclusion(self, item_id: int) -> BatchJobItem:
        """항목 수동 제외/포함 토글"""
        result = await self.db.execute(
            select(BatchJobItem).where(BatchJobItem.id == item_id)
        )
        item = result.scalar_one_or_none()
        if not item:
            raise ValueError("항목을 찾을 수 없습니다")

        item.manually_excluded = not item.manually_excluded
        item.updated_at = datetime.utcnow()
        await self.db.commit()
        return item

    async def get_item_detail(self, item_id: int) -> dict:
        """항목 상세: 수집내용(제개정이유/개정문) + AI 분석결과"""
        result = await self.db.execute(
            select(BatchJobItem).where(BatchJobItem.id == item_id)
        )
        item = result.scalar_one_or_none()
        if not item:
            raise ValueError("항목을 찾을 수 없습니다")

        # 수집내용: 해당 조례의 상위법령별 제개정이유
        mappings_result = await self.db.execute(
            select(OrdinanceLawMapping)
            .options(joinedload(OrdinanceLawMapping.law))
            .where(OrdinanceLawMapping.ordinance_id == item.ordinance_id)
        )
        mappings = list(mappings_result.scalars().all())

        revision_reasons = []
        for mapping in mappings:
            if not mapping.law:
                continue
            reason_result = await self.db.execute(
                select(LawRevisionReason).where(LawRevisionReason.law_id == mapping.law.id)
            )
            reason = reason_result.scalar_one_or_none()
            if reason:
                revision_reasons.append({
                    "law_id": mapping.law.id,
                    "law_name": mapping.law.law_name,
                    "revision_reason": reason.revision_reason,
                    "amendment_content": reason.amendment_content,
                    "extracted_articles": reason.extracted_articles,
                    "fetched_at": reason.fetched_at.isoformat() if reason.fetched_at else None,
                })

        # AI 분석결과: 해당 조례의 LLM 분석 결과
        ai_results_data = []
        for mapping in mappings:
            if not mapping.law:
                continue
            ai_result = await self.db.execute(
                select(LlmAnalysisResult).where(
                    LlmAnalysisResult.ordinance_id == item.ordinance_id,
                    LlmAnalysisResult.law_id == mapping.law.id,
                )
            )
            ai = ai_result.scalar_one_or_none()
            if ai:
                ai_results_data.append({
                    "law_id": mapping.law.id,
                    "law_name": mapping.law.law_name,
                    "status": ai.status,
                    "summary_text": ai.summary_text,
                    "review_draft_text": ai.review_draft_text,
                    "review_draft_result": ai.review_draft_result,
                    "affected_articles_json": ai.affected_articles_json,
                    "provider_name": ai.provider_name,
                    "model_name": ai.model_name,
                    "created_at": ai.created_at.isoformat() if ai.created_at else None,
                })

        return {
            "item": {
                "id": item.id,
                "ordinance_id": item.ordinance_id,
                "ordinance_name": item.ordinance_name,
                "ordinance_department": item.ordinance_department,
                "step2_result": item.step2_result,
                "step3_result": item.step3_result,
                "step4_result": item.step4_result,
                "step4_ai_result": item.step4_ai_result,
                "final_result": item.final_result,
            },
            "revision_reasons": revision_reasons,
            "ai_results": ai_results_data,
        }

    # ===== Step 2: 개정판별 =====

    async def run_step2_detect(self, job_id: int) -> AsyncGenerator[dict, None]:
        """Step 2: 대상 조례에 대해 개정 판별 실행 (SSE 스트리밍)"""
        job = await self.get_batch_job(job_id)
        if not job:
            raise ValueError("배치 작업을 찾을 수 없습니다")

        # 대상: step1 included + 수동 미제외 + step2 미처리
        items_result = await self.db.execute(
            select(BatchJobItem).where(
                BatchJobItem.batch_job_id == job_id,
                BatchJobItem.step1_result == "included",
                BatchJobItem.manually_excluded == False,
                BatchJobItem.step2_result.is_(None),
            )
        )
        items = list(items_result.scalars().all())

        job.current_step = "step2_detect"
        job.step2_status = "running"
        job.step2_total = len(items)
        job.step2_progress = 0
        await self.db.commit()

        detection_service = RevisionDetectionService(self.db)

        for i, item in enumerate(items):
            try:
                result = await detection_service.detect_all(item.ordinance_id)
                needs_revision = any(
                    r.get("needs_revision") for r in result.get("results", [])
                )

                if needs_revision:
                    item.step2_result = "needs_revision"
                    item.step2_reason = "개정 검토 필요"
                else:
                    item.step2_result = "no_revision"
                    item.step2_reason = "변경 해당없음"
                    item.final_result = "해당없음"

                item.step2_detail = {
                    "results": [
                        {
                            "method": r.get("method"),
                            "needs_revision": r.get("needs_revision"),
                        }
                        for r in result.get("results", [])
                    ]
                }

            except Exception as e:
                item.step2_result = "error"
                item.step2_reason = str(e)[:500]
                item.final_result = "오류(판별)"
                logger.exception(f"Step2 error for ordinance {item.ordinance_id}: {e}")

            item.updated_at = datetime.utcnow()
            job.step2_progress = i + 1
            await self.db.commit()

            yield {
                "progress": i + 1,
                "total": len(items),
                "ordinance_name": item.ordinance_name,
                "result": item.step2_result,
            }

        job.step2_status = "completed"
        job.current_step = "step2_detect"
        await self.db.commit()

    # ===== Step 3: 제개정이유 수집 =====

    async def run_step3_collect(self, job_id: int) -> AsyncGenerator[dict, None]:
        """Step 3: 개정대상 조례의 상위법령 제개정이유 수집"""
        job = await self.get_batch_job(job_id)
        if not job:
            raise ValueError("배치 작업을 찾을 수 없습니다")

        # 대상: step2 needs_revision + 수동 미제외 + step3 미처리
        items_result = await self.db.execute(
            select(BatchJobItem).where(
                BatchJobItem.batch_job_id == job_id,
                BatchJobItem.step2_result == "needs_revision",
                BatchJobItem.manually_excluded == False,
                BatchJobItem.step3_result.is_(None),
            )
        )
        items = list(items_result.scalars().all())

        job.current_step = "step3_collect"
        job.step3_status = "running"
        job.step3_total = len(items)
        job.step3_progress = 0
        await self.db.commit()

        from backend.external.moleg_client import MolegClient
        moleg_client = MolegClient(
            api_key=settings.MOLEG_API_KEY or "test",
            base_url=settings.MOLEG_API_BASE_URL,
        )

        try:
            for i, item in enumerate(items):
                try:
                    # 해당 조례의 상위법령 목록 조회
                    mappings_result = await self.db.execute(
                        select(OrdinanceLawMapping)
                        .options(joinedload(OrdinanceLawMapping.law))
                        .where(OrdinanceLawMapping.ordinance_id == item.ordinance_id)
                    )
                    mappings = list(mappings_result.scalars().all())

                    all_collected = True
                    for mapping in mappings:
                        if not mapping.law:
                            continue
                        # 이미 수집된 데이터 확인
                        existing = await self.db.execute(
                            select(LawRevisionReason).where(
                                LawRevisionReason.law_id == mapping.law.id
                            )
                        )
                        if existing.scalar_one_or_none():
                            continue

                        # 법제처 API로 제개정이유 수집
                        try:
                            detection_svc = RevisionDetectionService(self.db, moleg_client)
                            await detection_svc._get_or_fetch_revision_reason(mapping.law)
                        except Exception as e:
                            logger.warning(f"제개정이유 수집 실패 (law_id={mapping.law.id}): {e}")
                            all_collected = False

                    # 최종 수집 결과 확인 (row 존재 + 내용이 실제로 있는지)
                    has_any_data = False
                    for mapping in mappings:
                        if not mapping.law:
                            continue
                        check = await self.db.execute(
                            select(LawRevisionReason).where(
                                LawRevisionReason.law_id == mapping.law.id
                            )
                        )
                        reason_row = check.scalar_one_or_none()
                        if reason_row and (reason_row.revision_reason or reason_row.amendment_content):
                            has_any_data = True
                            break

                    if has_any_data:
                        item.step3_result = "collected"
                        item.step3_reason = "데이터 확보"
                    else:
                        item.step3_result = "no_data"
                        item.step3_reason = "제개정이유 데이터 없음"
                        item.final_result = "수동확인필요"

                except Exception as e:
                    item.step3_result = "error"
                    item.step3_reason = str(e)[:500]
                    item.final_result = "오류(수집)"
                    logger.exception(f"Step3 error for ordinance {item.ordinance_id}: {e}")

                item.updated_at = datetime.utcnow()
                job.step3_progress = i + 1
                await self.db.commit()

                yield {
                    "progress": i + 1,
                    "total": len(items),
                    "ordinance_name": item.ordinance_name,
                    "result": item.step3_result,
                }
        finally:
            await moleg_client.close()

        job.step3_status = "completed"
        job.current_step = "step3_collect"
        await self.db.commit()

    # ===== Step 4: AI 분석 =====

    async def run_step4_analyze(self, job_id: int) -> AsyncGenerator[dict, None]:
        """Step 4: 데이터 확보된 조례에 대해 AI 분석 실행"""
        job = await self.get_batch_job(job_id)
        if not job:
            raise ValueError("배치 작업을 찾을 수 없습니다")

        # 대상: step3 collected + 수동 미제외 + step4 미처리(이미 완료된 건 스킵)
        items_result = await self.db.execute(
            select(BatchJobItem).where(
                BatchJobItem.batch_job_id == job_id,
                BatchJobItem.step3_result == "collected",
                BatchJobItem.manually_excluded == False,
                BatchJobItem.step4_result.is_(None),  # 이미 처리된 건 스킵
            )
        )
        items = list(items_result.scalars().all())

        job.current_step = "step4_analyze"
        job.step4_status = "running"
        job.step4_total = len(items)
        job.step4_progress = 0
        await self.db.commit()

        analysis_service = LlmAnalysisService(self.db)

        for i, item in enumerate(items):
            try:
                # 해당 조례의 상위법령 조회
                mappings_result = await self.db.execute(
                    select(OrdinanceLawMapping)
                    .options(joinedload(OrdinanceLawMapping.law))
                    .where(OrdinanceLawMapping.ordinance_id == item.ordinance_id)
                )
                mappings = list(mappings_result.scalars().all())

                ai_results = []
                no_data_errors = []
                for mapping in mappings:
                    if not mapping.law:
                        continue
                    try:
                        result = await analysis_service.analyze_ordinance(
                            item.ordinance_id, mapping.law.id
                        )
                        ai_results.append(result)
                    except ConflictError:
                        # 이미 분석 완료 — 기존 결과 사용
                        existing = await self.db.execute(
                            select(LlmAnalysisResult).where(
                                LlmAnalysisResult.ordinance_id == item.ordinance_id,
                                LlmAnalysisResult.law_id == mapping.law.id,
                                LlmAnalysisResult.status == "success",
                            )
                        )
                        existing_result = existing.scalar_one_or_none()
                        if existing_result:
                            ai_results.append(existing_result)
                    except ValueError as e:
                        # 제개정이유 데이터 없음 등 데이터 부족 → 재시도해도 동일
                        no_data_errors.append(str(e))
                        logger.warning(f"AI 분석 데이터 부족 (ord={item.ordinance_id}, law={mapping.law.id}): {e}")
                    except Exception as e:
                        logger.warning(f"AI 분석 실패 (ord={item.ordinance_id}, law={mapping.law.id}): {e}")

                # 결과 집계
                if ai_results:
                    any_needs = any(
                        getattr(r, "review_draft_result", "") == "개정필요"
                        for r in ai_results
                    )
                    summaries = []
                    for r in ai_results:
                        if hasattr(r, "summary_text") and r.summary_text:
                            summaries.append(r.summary_text)

                    if any_needs:
                        item.step4_result = "needs_revision"
                        item.step4_ai_result = "개정필요"
                        item.final_result = "개정필요"
                    else:
                        item.step4_result = "no_revision"
                        item.step4_ai_result = "개정불필요"
                        item.final_result = "개정불필요"

                    item.step4_ai_summary = "\n\n---\n\n".join(summaries) if summaries else None
                    item.step4_reason = f"AI 분석 완료 ({len(ai_results)}건)"
                elif no_data_errors:
                    # 제개정이유 데이터 없음 → 수동확인 필요 (재시도 불필요)
                    item.step4_result = "no_revision"
                    item.step4_reason = "제개정이유 데이터 없음 — 수동확인 필요"
                    item.final_result = "수동확인필요"
                else:
                    item.step4_result = "error"
                    item.step4_reason = "AI 분석 결과 없음"
                    item.final_result = "오류(AI)"

            except Exception as e:
                item.step4_result = "error"
                item.step4_reason = str(e)[:500]
                item.final_result = "오류(AI)"
                logger.exception(f"Step4 error for ordinance {item.ordinance_id}: {e}")

            item.updated_at = datetime.utcnow()
            job.step4_progress = i + 1
            await self.db.commit()

            yield {
                "progress": i + 1,
                "total": len(items),
                "ordinance_name": item.ordinance_name,
                "result": item.step4_result,
            }

        job.step4_status = "completed"
        job.current_step = "step4_analyze"
        await self.db.commit()

    # ===== Step 5: 보고서 생성 =====

    async def generate_report_data(self, job_id: int) -> dict:
        """Step 5: 보고서용 데이터 생성"""
        job = await self.get_batch_job(job_id)
        if not job:
            raise ValueError("배치 작업을 찾을 수 없습니다")

        items_result = await self.db.execute(
            select(BatchJobItem)
            .where(BatchJobItem.batch_job_id == job_id)
            .order_by(BatchJobItem.ordinance_name)
        )
        items = list(items_result.scalars().all())

        # 통계 집계
        total = len(items)
        step1_included = sum(1 for i in items if i.step1_result == "included")
        step1_excluded = sum(1 for i in items if i.step1_result == "excluded")
        step2_needs = sum(1 for i in items if i.step2_result == "needs_revision")
        step2_no = sum(1 for i in items if i.step2_result == "no_revision")
        step2_error = sum(1 for i in items if i.step2_result == "error")
        step3_collected = sum(1 for i in items if i.step3_result == "collected")
        step3_no_data = sum(1 for i in items if i.step3_result == "no_data")
        step4_needs = sum(1 for i in items if i.step4_result == "needs_revision")
        step4_no = sum(1 for i in items if i.step4_result == "no_revision")

        # 부서별 집계
        dept_stats = {}
        for item in items:
            dept = item.ordinance_department or "미분류"
            if dept not in dept_stats:
                dept_stats[dept] = {"total": 0, "needs_revision": 0, "no_revision": 0, "excluded": 0, "error": 0}
            dept_stats[dept]["total"] += 1
            if item.final_result == "개정필요":
                dept_stats[dept]["needs_revision"] += 1
            elif item.final_result == "개정불필요":
                dept_stats[dept]["no_revision"] += 1
            elif item.final_result and item.final_result.startswith("제외"):
                dept_stats[dept]["excluded"] += 1
            elif item.final_result and "오류" in item.final_result:
                dept_stats[dept]["error"] += 1

        # 분류별 집계
        category_stats = {}
        for item in items:
            cat = item.ordinance_category or "미분류"
            if cat not in category_stats:
                category_stats[cat] = {"total": 0, "needs_revision": 0, "no_revision": 0}
            category_stats[cat]["total"] += 1
            if item.final_result == "개정필요":
                category_stats[cat]["needs_revision"] += 1
            elif item.final_result == "개정불필요":
                category_stats[cat]["no_revision"] += 1

        summary = {
            "job_name": job.name,
            "created_at": job.created_at.isoformat() if job.created_at else None,
            "total_ordinances": total,
            "pipeline": {
                "step1_included": step1_included,
                "step1_excluded": step1_excluded,
                "step2_needs_revision": step2_needs,
                "step2_no_revision": step2_no,
                "step2_error": step2_error,
                "step3_collected": step3_collected,
                "step3_no_data": step3_no_data,
                "step4_needs_revision": step4_needs,
                "step4_no_revision": step4_no,
            },
            "by_department": dept_stats,
            "by_category": category_stats,
        }

        # job에 요약 저장
        job.summary = summary
        job.step5_status = "completed"
        job.current_step = "step5_report"
        await self.db.commit()

        # 항목 상세 데이터 + 수집/분석 내용
        items_data = []
        for item in items:
            item_data = {
                "ordinance_id": item.ordinance_id,
                "ordinance_name": item.ordinance_name,
                "department": item.ordinance_department,
                "category": item.ordinance_category,
                "step1": item.step1_result,
                "step1_reason": item.step1_reason,
                "step2": item.step2_result,
                "step2_reason": item.step2_reason,
                "step3": item.step3_result,
                "step3_reason": item.step3_reason,
                "step4": item.step4_result,
                "step4_reason": item.step4_reason,
                "ai_result": item.step4_ai_result,
                "ai_summary": item.step4_ai_summary,
                "final_result": item.final_result,
                "revision_reasons": [],
                "ai_analysis": [],
            }

            # 수집내용 + AI 분석결과 조회 (개정대상 항목만)
            if item.step2_result == "needs_revision":
                mappings_result = await self.db.execute(
                    select(OrdinanceLawMapping)
                    .options(joinedload(OrdinanceLawMapping.law))
                    .where(OrdinanceLawMapping.ordinance_id == item.ordinance_id)
                )
                mappings = list(mappings_result.scalars().all())

                for mapping in mappings:
                    if not mapping.law:
                        continue
                    # 수집내용
                    rr = await self.db.execute(
                        select(LawRevisionReason).where(LawRevisionReason.law_id == mapping.law.id)
                    )
                    reason = rr.scalar_one_or_none()
                    if reason:
                        item_data["revision_reasons"].append({
                            "law_name": mapping.law.law_name,
                            "revision_reason": reason.revision_reason,
                            "amendment_content": reason.amendment_content,
                        })

                    # AI 분석결과
                    ar = await self.db.execute(
                        select(LlmAnalysisResult).where(
                            LlmAnalysisResult.ordinance_id == item.ordinance_id,
                            LlmAnalysisResult.law_id == mapping.law.id,
                            LlmAnalysisResult.status == "success",
                        )
                    )
                    ai = ar.scalar_one_or_none()
                    if ai:
                        item_data["ai_analysis"].append({
                            "law_name": mapping.law.law_name,
                            "summary_text": ai.summary_text,
                            "review_draft_text": ai.review_draft_text,
                            "review_draft_result": ai.review_draft_result,
                            "affected_articles_json": ai.affected_articles_json,
                            "provider_name": ai.provider_name,
                            "model_name": ai.model_name,
                        })

            items_data.append(item_data)

        return {"summary": summary, "items": items_data}

    async def export_excel(self, job_id: int) -> bytes:
        """Excel 보고서 생성"""
        import io
        import pandas as pd

        report = await self.generate_report_data(job_id)
        items = report["items"]
        summary = report["summary"]

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            # Sheet 1: 요약
            summary_rows = [
                ["작업명", summary["job_name"]],
                ["생성일시", summary["created_at"]],
                ["전체 조례 수", summary["total_ordinances"]],
                ["", ""],
                ["파이프라인 요약", ""],
                ["Step 1 대상", summary["pipeline"]["step1_included"]],
                ["Step 1 제외", summary["pipeline"]["step1_excluded"]],
                ["Step 2 개정대상", summary["pipeline"]["step2_needs_revision"]],
                ["Step 2 해당없음", summary["pipeline"]["step2_no_revision"]],
                ["Step 2 오류", summary["pipeline"]["step2_error"]],
                ["Step 3 데이터확보", summary["pipeline"]["step3_collected"]],
                ["Step 3 데이터없음", summary["pipeline"]["step3_no_data"]],
                ["Step 4 개정필요", summary["pipeline"]["step4_needs_revision"]],
                ["Step 4 개정불필요", summary["pipeline"]["step4_no_revision"]],
            ]
            df_summary = pd.DataFrame(summary_rows, columns=["항목", "값"])
            df_summary.to_excel(writer, sheet_name="요약", index=False)

            # Sheet 2: 부서별 집계
            dept_rows = []
            for dept, stats in summary.get("by_department", {}).items():
                dept_rows.append({
                    "부서": dept,
                    "전체": stats["total"],
                    "개정필요": stats["needs_revision"],
                    "개정불필요": stats["no_revision"],
                    "제외": stats["excluded"],
                    "오류": stats["error"],
                })
            if dept_rows:
                df_dept = pd.DataFrame(dept_rows)
                df_dept.to_excel(writer, sheet_name="부서별", index=False)

            # Sheet 3: 전체 상세 목록
            detail_rows = []
            for item in items:
                detail_rows.append({
                    "조례명": item["ordinance_name"],
                    "소관부서": item["department"],
                    "분류": item["category"],
                    "Step1(대상선정)": item["step1"],
                    "Step1 사유": item["step1_reason"] or "",
                    "Step2(개정판별)": item["step2"] or "",
                    "Step2 사유": item["step2_reason"] or "",
                    "Step3(데이터수집)": item["step3"] or "",
                    "Step3 사유": item["step3_reason"] or "",
                    "Step4(AI분석)": item["step4"] or "",
                    "AI 판정": item["ai_result"] or "",
                    "최종결과": item["final_result"] or "",
                })
            df_detail = pd.DataFrame(detail_rows)
            df_detail.to_excel(writer, sheet_name="상세목록", index=False)

            # Sheet 4: 개정필요 목록 (수집내용 + AI 분석 포함)
            revision_rows = []
            for item in items:
                if item["final_result"] == "개정필요":
                    # 수집내용 텍스트
                    reasons_text = ""
                    for rr in item.get("revision_reasons", []):
                        reasons_text += f"[{rr['law_name']}]\n"
                        if rr.get("revision_reason"):
                            reasons_text += f"제개정이유: {rr['revision_reason'][:500]}\n"
                        if rr.get("amendment_content"):
                            reasons_text += f"개정문: {rr['amendment_content'][:500]}\n"
                        reasons_text += "\n"

                    # AI 분석 텍스트
                    ai_text = ""
                    for ai in item.get("ai_analysis", []):
                        ai_text += f"[{ai['law_name']}] ({ai.get('review_draft_result', '')})\n"
                        if ai.get("summary_text"):
                            ai_text += f"{ai['summary_text']}\n"
                        if ai.get("review_draft_text"):
                            ai_text += f"검토의견: {ai['review_draft_text']}\n"
                        ai_text += "\n"

                    revision_rows.append({
                        "조례명": item["ordinance_name"],
                        "소관부서": item["department"],
                        "분류": item["category"],
                        "수집내용(제개정이유)": reasons_text.strip(),
                        "AI 분석 결과": ai_text.strip() or item.get("ai_summary") or "",
                    })
            if revision_rows:
                df_revision = pd.DataFrame(revision_rows)
                df_revision.to_excel(writer, sheet_name="개정필요목록", index=False)

            # Sheet 5: 개정불필요 목록
            no_revision_rows = []
            for item in items:
                if item["final_result"] == "개정불필요":
                    ai_text = ""
                    for ai in item.get("ai_analysis", []):
                        ai_text += f"[{ai['law_name']}] {ai.get('review_draft_result', '')}\n"
                        if ai.get("summary_text"):
                            ai_text += f"{ai['summary_text']}\n"
                    no_revision_rows.append({
                        "조례명": item["ordinance_name"],
                        "소관부서": item["department"],
                        "분류": item["category"],
                        "AI 분석 결과": ai_text.strip() or item.get("ai_summary") or "",
                    })
            if no_revision_rows:
                df_no = pd.DataFrame(no_revision_rows)
                df_no.to_excel(writer, sheet_name="개정불필요목록", index=False)

        return output.getvalue()

    async def export_pdf(self, job_id: int) -> bytes:
        """PDF 보고서 생성 — 완성된 문서서식, 분석 건별 챕터"""
        import io
        from pathlib import Path
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table as RLTable,
            TableStyle, PageBreak, HRFlowable,
        )
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        report = await self.generate_report_data(job_id)
        summary = report["summary"]
        items = report["items"]

        output = io.BytesIO()
        doc = SimpleDocTemplate(
            output, pagesize=A4,
            leftMargin=20*mm, rightMargin=20*mm,
            topMargin=20*mm, bottomMargin=20*mm,
        )

        # 한글 폰트 등록 — Docker(/app/fonts) + Windows 경로
        font_name = "Helvetica"
        for font_path in [
            "/app/fonts/malgun.ttf",
            "C:/Windows/Fonts/malgun.ttf",
            "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        ]:
            try:
                if Path(font_path).exists():
                    pdfmetrics.registerFont(TTFont("KoreanFont", font_path))
                    font_name = "KoreanFont"
                    break
            except Exception:
                continue

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("KTitle", parent=styles["Title"], fontName=font_name, fontSize=20, spaceAfter=12)
        h1_style = ParagraphStyle("KH1", parent=styles["Heading1"], fontName=font_name, fontSize=16, spaceBefore=20, spaceAfter=10)
        h2_style = ParagraphStyle("KH2", parent=styles["Heading2"], fontName=font_name, fontSize=13, spaceBefore=14, spaceAfter=6)
        h3_style = ParagraphStyle("KH3", parent=styles["Heading3"], fontName=font_name, fontSize=11, spaceBefore=8, spaceAfter=4)
        body_style = ParagraphStyle("KBody", parent=styles["Normal"], fontName=font_name, fontSize=9, leading=13)
        small_style = ParagraphStyle("KSmall", parent=styles["Normal"], fontName=font_name, fontSize=8, leading=11, textColor=colors.grey)
        label_style = ParagraphStyle("KLabel", parent=styles["Normal"], fontName=font_name, fontSize=9, leading=12, textColor=colors.HexColor("#4472C4"))

        def _esc(text: str, limit: int = 5000) -> str:
            """XML 이스케이프 + 줄바꿈 변환"""
            if not text:
                return ""
            t = text[:limit].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            return t.replace("\n", "<br/>")

        def _make_table(data, col_widths):
            t = RLTable(data, colWidths=col_widths, repeatRows=1)
            t.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8F8F8")]),
            ]))
            return t

        elements = []

        # ============ 표지 ============
        elements.append(Spacer(1, 50*mm))
        elements.append(Paragraph("일괄 개정검토 보고서", title_style))
        elements.append(Spacer(1, 15*mm))
        elements.append(Paragraph(f"작업명: {_esc(summary['job_name'])}", body_style))
        elements.append(Paragraph(f"생성일: {summary.get('created_at', '')}", body_style))
        elements.append(Paragraph(f"전체 대상 조례: {summary['total_ordinances']}건", body_style))
        pipe = summary["pipeline"]
        elements.append(Spacer(1, 5*mm))
        elements.append(Paragraph(f"개정필요: {pipe['step4_needs_revision']}건 / 개정불필요: {pipe['step4_no_revision']}건", body_style))
        elements.append(PageBreak())

        # ============ 제1장. 요약 ============
        elements.append(Paragraph("제1장 파이프라인 요약", h1_style))
        summary_data = [
            ["단계", "항목", "건수"],
            ["1. 대상 선정", "대상", str(pipe["step1_included"])],
            ["", "제외", str(pipe["step1_excluded"])],
            ["2. 개정 판별", "개정대상", str(pipe["step2_needs_revision"])],
            ["", "해당없음", str(pipe["step2_no_revision"])],
            ["3. 데이터 수집", "확보", str(pipe["step3_collected"])],
            ["", "미확보", str(pipe["step3_no_data"])],
            ["4. AI 분석", "개정필요", str(pipe["step4_needs_revision"])],
            ["", "개정불필요", str(pipe["step4_no_revision"])],
        ]
        elements.append(_make_table(summary_data, [55*mm, 45*mm, 30*mm]))
        elements.append(Spacer(1, 10*mm))

        # 부서별 집계
        elements.append(Paragraph("부서별 집계", h2_style))
        dept_data = [["부서", "전체", "개정필요", "개정불필요", "제외", "오류"]]
        for dept, stats in summary.get("by_department", {}).items():
            dept_data.append([dept, str(stats["total"]), str(stats["needs_revision"]),
                              str(stats["no_revision"]), str(stats["excluded"]), str(stats["error"])])
        if len(dept_data) > 1:
            elements.append(_make_table(dept_data, [50*mm, 20*mm, 20*mm, 22*mm, 20*mm, 20*mm]))
        elements.append(PageBreak())

        # ============ 제2장. 분석 대상 조례 상세 — 건별 챕터 ============
        analyzed_items = [i for i in items if i["final_result"] in ("개정필요", "개정불필요")]
        elements.append(Paragraph(f"제2장 분석 대상 조례 상세 ({len(analyzed_items)}건)", h1_style))

        if not analyzed_items:
            elements.append(Paragraph("분석 대상 조례가 없습니다.", body_style))
        else:
            for idx, item in enumerate(analyzed_items, 1):
                if idx > 1:
                    elements.append(PageBreak())

                # 조례 제목 + 판정
                result_label = item.get("ai_result") or item.get("final_result") or ""
                result_color = "#cf1322" if result_label == "개정필요" else "#389e0d"
                elements.append(Paragraph(
                    f"2-{idx}. {_esc(item['ordinance_name'])}",
                    h2_style,
                ))
                # 기본 정보 테이블
                info_data = [
                    ["소관부서", item["department"] or "-", "분류", item["category"] or "-"],
                    ["AI 판정", result_label, "", ""],
                ]
                info_t = RLTable(info_data, colWidths=[25*mm, 55*mm, 25*mm, 55*mm])
                info_t.setStyle(TableStyle([
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8EDF3")),
                    ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#E8EDF3")),
                ]))
                elements.append(info_t)
                elements.append(Spacer(1, 4*mm))

                # ---- 수집내용 ----
                revision_reasons = item.get("revision_reasons", [])
                if revision_reasons:
                    elements.append(Paragraph("수집내용 (제개정이유 / 개정문)", label_style))
                    for rr in revision_reasons:
                        elements.append(Paragraph(f"<b>상위법령: {_esc(rr['law_name'])}</b>", body_style))
                        if rr.get("revision_reason"):
                            elements.append(Paragraph(f"<b>제개정이유:</b><br/>{_esc(rr['revision_reason'], 2000)}", body_style))
                        if rr.get("amendment_content"):
                            elements.append(Paragraph(f"<b>개정문:</b><br/>{_esc(rr['amendment_content'], 2000)}", body_style))
                        elements.append(Spacer(1, 2*mm))

                # ---- AI 분석결과 ----
                ai_analysis = item.get("ai_analysis", [])
                if ai_analysis:
                    elements.append(Spacer(1, 3*mm))
                    elements.append(Paragraph("AI 분석결과", label_style))
                    for ai in ai_analysis:
                        elements.append(Paragraph(
                            f"<b>{_esc(ai['law_name'])} — {ai.get('review_draft_result', '')}</b> "
                            f"({ai.get('provider_name', '')} {ai.get('model_name', '')})",
                            body_style,
                        ))
                        if ai.get("summary_text"):
                            elements.append(Paragraph(_esc(ai["summary_text"], 3000), body_style))

                        # 영향받는 조문
                        affected = ai.get("affected_articles_json") or []
                        if affected:
                            elements.append(Spacer(1, 2*mm))
                            elements.append(Paragraph("<b>영향받는 조례 조문:</b>", body_style))
                            for art in affected:
                                art_text = f"  {_esc(art.get('article_no', ''))}"
                                if art.get("article_title"):
                                    art_text += f"({_esc(art['article_title'])})"
                                if art.get("issue"):
                                    art_text += f" — 문제: {_esc(art['issue'])}"
                                if art.get("recommendation"):
                                    art_text += f" / 권고: {_esc(art['recommendation'])}"
                                elements.append(Paragraph(art_text, body_style))

                        if ai.get("review_draft_text"):
                            elements.append(Spacer(1, 2*mm))
                            elements.append(Paragraph(f"<b>검토의견 초안:</b><br/>{_esc(ai['review_draft_text'], 2000)}", body_style))

                        elements.append(Spacer(1, 3*mm))

                if not revision_reasons and not ai_analysis:
                    elements.append(Paragraph("수집/분석 데이터가 없습니다.", small_style))

        elements.append(PageBreak())

        # ============ 제3장. 전체 현황 ============
        elements.append(Paragraph("제3장 전체 조례 현황", h1_style))
        all_data = [["No.", "조례명", "소관부서", "분류", "최종결과"]]
        for i, item in enumerate(items, 1):
            all_data.append([
                str(i),
                item["ordinance_name"][:35],
                item["department"] or "",
                item["category"] or "",
                item["final_result"] or "-",
            ])
        elements.append(_make_table(all_data, [12*mm, 55*mm, 32*mm, 18*mm, 40*mm]))

        doc.build(elements)
        return output.getvalue()

    # ===== Step 2/3/4 재시도 =====

    async def retry_step_items(
        self, job_id: int, step: str, item_ids: Optional[list[int]] = None
    ) -> int:
        """특정 단계의 오류 항목 재시도를 위해 상태 초기화"""
        query = select(BatchJobItem).where(
            BatchJobItem.batch_job_id == job_id,
            getattr(BatchJobItem, f"{step}_result") == "error",
        )
        if item_ids:
            query = query.where(BatchJobItem.id.in_(item_ids))

        result = await self.db.execute(query)
        items = list(result.scalars().all())

        for item in items:
            setattr(item, f"{step}_result", None)
            setattr(item, f"{step}_reason", None)
            item.final_result = None
            item.updated_at = datetime.utcnow()

        await self.db.commit()
        return len(items)

    async def export_docx(self, job_id: int) -> bytes:
        """Word(docx) 보고서 생성 — 건별 챕터"""
        import io
        from docx import Document
        from docx.shared import Pt, Cm, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        report = await self.generate_report_data(job_id)
        summary = report["summary"]
        items = report["items"]
        pipe = summary["pipeline"]

        doc = Document()

        # 한글 폰트 설정 — rFonts eastAsia 속성 필수
        from docx.oxml.ns import qn

        def _set_korean_font(style, font_name="맑은 고딕", size=Pt(10)):
            style.font.name = font_name
            style.font.size = size
            rpr = style.element.rPr
            if rpr is None:
                style.element.get_or_add_rPr()
                rpr = style.element.rPr
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                from lxml import etree
                rfonts = etree.SubElement(rpr, qn("w:rFonts"))
            rfonts.set(qn("w:eastAsia"), font_name)

        _set_korean_font(doc.styles["Normal"])
        for sname in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "No Spacing"]:
            try:
                _set_korean_font(doc.styles[sname])
            except KeyError:
                pass

        # ============ 표지 ============
        doc.add_paragraph("")
        doc.add_paragraph("")
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("일괄 개정검토 보고서")
        run.font.size = Pt(24)
        run.bold = True

        doc.add_paragraph("")
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"작업명: {summary['job_name']}").font.size = Pt(12)
        info2 = doc.add_paragraph()
        info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info2.add_run(f"생성일: {summary.get('created_at', '')}").font.size = Pt(11)
        info3 = doc.add_paragraph()
        info3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info3.add_run(
            f"전체 {summary['total_ordinances']}건 | "
            f"개정필요 {pipe['step4_needs_revision']}건 | "
            f"개정불필요 {pipe['step4_no_revision']}건"
        ).font.size = Pt(11)

        doc.add_page_break()

        # ============ 제1장 요약 ============
        doc.add_heading("제1장 파이프라인 요약", level=1)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "단계"
        hdr[1].text = "항목"
        hdr[2].text = "건수"
        summary_rows = [
            ("1. 대상 선정", "대상", pipe["step1_included"]),
            ("", "제외", pipe["step1_excluded"]),
            ("2. 개정 판별", "개정대상", pipe["step2_needs_revision"]),
            ("", "해당없음", pipe["step2_no_revision"]),
            ("3. 데이터 수집", "확보", pipe["step3_collected"]),
            ("", "미확보", pipe["step3_no_data"]),
            ("4. AI 분석", "개정필요", pipe["step4_needs_revision"]),
            ("", "개정불필요", pipe["step4_no_revision"]),
        ]
        for step_name, label, count in summary_rows:
            row = table.add_row().cells
            row[0].text = step_name
            row[1].text = label
            row[2].text = str(count)

        doc.add_paragraph("")
        doc.add_heading("부서별 집계", level=2)
        dept_items = list(summary.get("by_department", {}).items())
        if dept_items:
            dt = doc.add_table(rows=1, cols=6)
            dt.style = "Light Grid Accent 1"
            dt.alignment = WD_TABLE_ALIGNMENT.CENTER
            dh = dt.rows[0].cells
            for i, h in enumerate(["부서", "전체", "개정필요", "개정불필요", "제외", "오류"]):
                dh[i].text = h
            for dept, stats in dept_items:
                row = dt.add_row().cells
                row[0].text = dept
                row[1].text = str(stats["total"])
                row[2].text = str(stats["needs_revision"])
                row[3].text = str(stats["no_revision"])
                row[4].text = str(stats["excluded"])
                row[5].text = str(stats["error"])

        doc.add_page_break()

        # ============ 제2장 분석 대상 조례 상세 ============
        analyzed = [i for i in items if i["final_result"] in ("개정필요", "개정불필요")]
        doc.add_heading(f"제2장 분석 대상 조례 상세 ({len(analyzed)}건)", level=1)

        for idx, item in enumerate(analyzed, 1):
            if idx > 1:
                doc.add_page_break()

            # 조례 제목
            result_label = item.get("ai_result") or item.get("final_result") or ""
            heading = doc.add_heading(level=2)
            heading.add_run(f"2-{idx}. {item['ordinance_name']}")

            # 기본 정보
            p = doc.add_paragraph()
            p.add_run("소관부서: ").bold = True
            p.add_run(f"{item['department'] or '-'}    ")
            p.add_run("분류: ").bold = True
            p.add_run(f"{item['category'] or '-'}    ")
            p.add_run("AI 판정: ").bold = True
            result_run = p.add_run(result_label)
            if result_label == "개정필요":
                result_run.font.color.rgb = RGBColor(0xCF, 0x13, 0x22)
                result_run.bold = True
            elif result_label == "개정불필요":
                result_run.font.color.rgb = RGBColor(0x38, 0x9E, 0x0D)

            # 수집내용
            revision_reasons = item.get("revision_reasons", [])
            if revision_reasons:
                doc.add_heading("수집내용 (제개정이유 / 개정문)", level=3)
                for rr in revision_reasons:
                    p = doc.add_paragraph()
                    p.add_run(f"상위법령: {rr['law_name']}").bold = True

                    if rr.get("revision_reason"):
                        doc.add_paragraph("")
                        pr = doc.add_paragraph()
                        pr.add_run("제개정이유: ").bold = True
                        pr.add_run(rr["revision_reason"][:3000])

                    if rr.get("amendment_content"):
                        pa = doc.add_paragraph()
                        pa.add_run("개정문: ").bold = True
                        pa.add_run(rr["amendment_content"][:3000])

                    doc.add_paragraph("")

            # AI 분석결과
            ai_analysis = item.get("ai_analysis", [])
            if ai_analysis:
                doc.add_heading("AI 분석결과", level=3)
                for ai in ai_analysis:
                    p = doc.add_paragraph()
                    p.add_run(f"{ai['law_name']} — {ai.get('review_draft_result', '')}").bold = True
                    run_info = p.add_run(f"  ({ai.get('provider_name', '')} {ai.get('model_name', '')})")
                    run_info.font.size = Pt(8)
                    run_info.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

                    if ai.get("summary_text"):
                        doc.add_paragraph(ai["summary_text"][:5000])

                    affected = ai.get("affected_articles_json") or []
                    if affected:
                        doc.add_paragraph("")
                        pa = doc.add_paragraph()
                        pa.add_run("영향받는 조례 조문:").bold = True
                        for art in affected:
                            art_p = doc.add_paragraph(style="List Bullet")
                            art_text = art.get("article_no", "")
                            if art.get("article_title"):
                                art_text += f"({art['article_title']})"
                            if art.get("issue"):
                                art_text += f" — 문제: {art['issue']}"
                            if art.get("recommendation"):
                                art_text += f" / 권고: {art['recommendation']}"
                            art_p.add_run(art_text)

                    if ai.get("review_draft_text"):
                        doc.add_paragraph("")
                        rd = doc.add_paragraph()
                        rd.add_run("검토의견 초안: ").bold = True
                        rd.add_run(ai["review_draft_text"][:3000])

                    doc.add_paragraph("")

            if not revision_reasons and not ai_analysis:
                doc.add_paragraph("수집/분석 데이터가 없습니다.", style="No Spacing")

        doc.add_page_break()

        # ============ 제3장 전체 현황 ============
        doc.add_heading("제3장 전체 조례 현황", level=1)
        at = doc.add_table(rows=1, cols=5)
        at.style = "Light Grid Accent 1"
        at.alignment = WD_TABLE_ALIGNMENT.CENTER
        ah = at.rows[0].cells
        for i, h in enumerate(["No.", "조례명", "소관부서", "분류", "최종결과"]):
            ah[i].text = h
        for i, item in enumerate(items, 1):
            row = at.add_row().cells
            row[0].text = str(i)
            row[1].text = item["ordinance_name"][:40]
            row[2].text = item["department"] or ""
            row[3].text = item["category"] or ""
            row[4].text = item["final_result"] or "-"

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
